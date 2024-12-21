import os
from django.core.management.base import BaseCommand
from docx import Document

class Command(BaseCommand):
    help = 'Generate SQL file for bulk uploading MCQs from DOC files'

    VALID_DIFFICULTIES = ['Easy', 'Medium', 'Tough']
    VALID_TYPES = ['General', 'Clinical', 'Image']

    def handle(self, *args, **kwargs):
        directory = '/path/to/your/doc/files/'  # Change this to your directory with DOCX files
        output_file = '/path/to/output/sql/file.sql'  # SQL output file

        with open(output_file, 'w') as sql_file:
            for filename in os.listdir(directory):
                if filename.endswith('.docx'):
                    filepath = os.path.join(directory, filename)
                    self.process_doc(filepath, sql_file)

    def process_doc(self, filepath, sql_file):
        document = Document(filepath)
        current_subject = "Mathematics"  # Since one file represents one subject, set it here
        current_unit = None
        current_chapter = None
        current_topic = None

        # Insert Subject
        subject_sql = f"""
        INSERT INTO yourapp_subject (name) 
        VALUES ('{self.escape_quotes(current_subject)}')
        ON CONFLICT (name) DO NOTHING;
        """
        sql_file.write(subject_sql)

        for para in document.paragraphs:
            line = para.text.strip()

            if not line:
                continue

            if line.startswith('Unit:'):
                unit_name = line.replace('Unit:', '').strip()
                unit_sql = f"""
                INSERT INTO yourapp_unit (name, subject_id) 
                VALUES ('{self.escape_quotes(unit_name)}', 
                (SELECT id FROM yourapp_subject WHERE name='{self.escape_quotes(current_subject)}'))
                ON CONFLICT (name) DO NOTHING;
                """
                sql_file.write(unit_sql)
                current_unit = unit_name

            elif line.startswith('Chapter:'):
                chapter_name = line.replace('Chapter:', '').strip()
                chapter_sql = f"""
                INSERT INTO yourapp_chapter (name, unit_id) 
                VALUES ('{self.escape_quotes(chapter_name)}', 
                (SELECT id FROM yourapp_unit WHERE name='{self.escape_quotes(current_unit)}'))
                ON CONFLICT (name) DO NOTHING;
                """
                sql_file.write(chapter_sql)
                current_chapter = chapter_name

            elif line.startswith('Topic:'):
                topic_name = line.replace('Topic:', '').strip()
                topic_sql = f"""
                INSERT INTO yourapp_topic (name, chapter_id) 
                VALUES ('{self.escape_quotes(topic_name)}', 
                (SELECT id FROM yourapp_chapter WHERE name='{self.escape_quotes(current_chapter)}'))
                ON CONFLICT (name) DO NOTHING;
                """
                sql_file.write(topic_sql)
                current_topic = topic_name

            elif line.startswith('MCQ'):
                mcq_text = line.split(':', 1)[1].strip()
                options = []
                correct_option = None
                explanation = None
                difficulty = None
                mcq_type = None

                # Collect the options (next 4 lines after MCQ)
                for _ in range(4):
                    option = next(document.paragraphs).text.strip().split(". ", 1)[1]
                    options.append(self.escape_quotes(option))

                # Read extra fields like Correct, Explanation, Difficulty, Type
                for next_line in document.paragraphs:
                    next_text = next_line.text.strip()
                    if next_text.startswith('Correct:'):
                        correct_option = self.escape_quotes(next_text.replace('Correct:', '').strip())
                    elif next_text.startswith('Explanation:'):
                        explanation = self.escape_quotes(next_text.replace('Explanation:', '').strip())
                    elif next_text.startswith('Difficulty:'):
                        difficulty = self.escape_quotes(next_text.replace('Difficulty:', '').strip())
                        if difficulty not in self.VALID_DIFFICULTIES:
                            raise ValueError(f"Invalid difficulty: {difficulty}")
                    elif next_text.startswith('Type:'):
                        mcq_type = self.escape_quotes(next_text.replace('Type:', '').strip())
                        if mcq_type not in self.VALID_TYPES:
                            raise ValueError(f"Invalid MCQ type: {mcq_type}")
                    elif not next_text:  # Stop when a blank line is encountered
                        break

                # Create SQL insert statement for MCQ
                mcq_sql = f"""
                INSERT INTO yourapp_mcq (text, option_1, option_2, option_3, option_4, correct_option, explanation, topic_id, difficulty_id, types_id) 
                VALUES (
                    '{self.escape_quotes(mcq_text)}', 
                    '{options[0]}', 
                    '{options[1]}', 
                    '{options[2]}', 
                    '{options[3]}', 
                    '{correct_option}', 
                    '{explanation}', 
                    (SELECT id FROM yourapp_topic WHERE name='{self.escape_quotes(current_topic)}'),
                    (SELECT id FROM yourapp_difficulties WHERE name='{self.escape_quotes(difficulty)}'), 
                    (SELECT id FROM yourapp_mcq_types WHERE types='{self.escape_quotes(mcq_type)}')
                );
                """
                sql_file.write(mcq_sql)

        self.stdout.write(self.style.SUCCESS(f"Successfully generated SQL file for {os.path.basename(filepath)}"))

    def escape_quotes(self, text):
        """Helper function to escape single quotes for PostgreSQL compatibility."""
        return text.replace("'", "''")

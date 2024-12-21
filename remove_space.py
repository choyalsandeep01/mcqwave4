import docx
import re

def convert_mcq_to_single_line(paragraphs):
    """
    Convert a multi-line MCQ to a single-line format
    
    Args:
    paragraphs (list): List of paragraph texts related to a single MCQ
    
    Returns:
    str: Single-line MCQ text
    """
    # Join all paragraphs into a single string
    full_text = ' '.join([p.strip() for p in paragraphs])
    
    # Remove extra spaces around '|'
    full_text = re.sub(r'\s*\|\s*', '|', full_text)
    
    return full_text

def process_mcq_file(input_file, output_file):
    """
    Process DOCX file to convert multi-line MCQs to single-line format
    
    Args:
    input_file (str): Path to input DOCX file
    output_file (str): Path to output DOCX file
    """
    # Open the input document
    doc = docx.Document(input_file)
    
    # Create a new document
    output_doc = docx.Document()
    
    # Variables to track MCQ construction
    current_mcq_paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        
        # Skip completely empty paragraphs
        if not text:
            continue
        
        # Add to current MCQ paragraphs
        current_mcq_paragraphs.append(text)
        
        # Check if this is the end of an MCQ (contains difficulty and category)
        if re.search(r'\|[A-Za-z]+\|[A-Za-z]+\|$', text):
            # Convert multi-line MCQ to single line
            single_line_mcq = convert_mcq_to_single_line(current_mcq_paragraphs)
            
            # Add to output document
            output_para = output_doc.add_paragraph(single_line_mcq)
            
            # Reset MCQ paragraphs
            current_mcq_paragraphs = []
    
    # Save the output document
    output_doc.save(output_file)
    print(f"Processed MCQs saved to {output_file}")

# Example usage
if __name__ == "__main__":
    input_file = "IMMUNOLOGY 1A.docx"  # Replace with your input file path
    output_file = "single_line_mcqs.docx"  # Replace with your desired output file path
    process_mcq_file(input_file, output_file)

# Note: Requires python-docx library
# Install with: pip install python-docx
"""
Key Features:
1. Converts multi-line MCQs to single-line format
2. Removes extra spaces around '|'
3. Preserves full MCQ content
4. Works with varying MCQ structures
5. Handles multiple MCQs in a document
"""
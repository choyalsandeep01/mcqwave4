import os
import django
import traceback
import logging
import docx
import re

# Set up Django environment
os.environ.setdefault('DJANGO_SETTINGS_MODULE', 'myproject.settings')
django.setup()

# Import your models
from mcqs.models import Subject, Unit, Chapter, Topic, difficulties, mcq_types, MCQ

# Configure logging
logging.basicConfig(level=logging.DEBUG, 
                    format='%(asctime)s - %(levelname)s - %(message)s',
                    filename='mcq_upload_log.txt')
logger = logging.getLogger(__name__)

def normalize_text(text):
    """
    Normalize text to match database entries
    Converts to title case, ensuring first letter is capital
    """
    return text.strip().title()

def validate_difficulty(difficulty_name):
    """
    Validate if difficulty exists in the database
    """
    try:
        normalized_difficulty = normalize_text(difficulty_name)
        difficulty = difficulties.objects.filter(name=normalized_difficulty).first()
        
        if not difficulty:
            logger.error(f"Difficulty '{difficulty_name}' does not exist in the database.")
            return None
        
        return difficulty
    except Exception as e:
        logger.error(f"Error validating difficulty {difficulty_name}: {e}")
        return None

def validate_mcq_type(mcq_type):
    """
    Validate if MCQ type exists in the database
    """
    try:
        normalized_type = normalize_text(mcq_type)
        mcq_type_obj = mcq_types.objects.filter(types=normalized_type).first()
        
        if not mcq_type_obj:
            logger.error(f"MCQ Type '{mcq_type}' does not exist in the database.")
            return None
        
        return mcq_type_obj
    except Exception as e:
        logger.error(f"Error validating MCQ type {mcq_type}: {e}")
        return None

def parse_mcq(paragraphs):
    """
    Parse MCQ entry with enhanced flexibility
    Handles both single-line and multi-line formats
    """
    # Join all paragraphs, removing extra whitespaces
    full_text = ' '.join([p.text.strip() for p in paragraphs])
    full_text = re.sub(r'\s+', ' ', full_text)
    
    logger.debug(f"Full MCQ text: {full_text}")
    
    try:
        # First try to split by '|'
        parts = [p.strip() for p in full_text.split('|')]
        
        # Validate minimum number of parts
        if len(parts) < 7:
            # If splitting fails, try alternative parsing
            logger.error(f"Insufficient parts in MCQ: {full_text}")
            return None
        
        # Extract question (first part)
        question = parts[0]
        
        # Handling options and correct answer
        if len(parts) >= 7:
            options = parts[1:5]
            correct_answer = parts[5]
            explanation = parts[6]
            
            # Default difficulty and type if not specified
            difficulty = parts[7] if len(parts) > 7 else 'Easy'
            mcq_type = parts[8] if len(parts) > 8 else 'General'
            
            # Construct full parts list
            full_parts = [
                question,      # 0: Question
                options[0],    # 1: Option 1
                options[1],    # 2: Option 2
                options[2],    # 3: Option 3
                options[3],    # 4: Option 4
                correct_answer,# 5: Correct Option
                explanation,   # 6: Explanation
                difficulty,    # 7: Difficulty
                mcq_type       # 8: MCQ Type
            ]
            
            logger.debug(f"Parsed MCQ parts: {full_parts}")
            return full_parts
        
        return None
    
    except Exception as e:
        logger.error(f"Error parsing MCQ: {e}")
        logger.error(f"Full text: {full_text}")
        return None

def process_docx_file(file_path):
    """
    Process DOCX file and upload to database
    """
    try:
        # Open the Word document
        doc = docx.Document(file_path)
        
        # Get subject name from filename
        subject_name = os.path.splitext(os.path.basename(file_path))[0]
        
        # Create or get subject
        subject, _ = Subject.objects.get_or_create(name=subject_name)
        
        current_unit = None
        current_chapter = None
        current_topic = None
        
        # Track successfully uploaded MCQs to remove from document
        uploaded_mcq_indices = []
        
        # Iterate through paragraphs
        i = 0
        while i < len(doc.paragraphs):
            current_para = doc.paragraphs[i]
            text = current_para.text.strip()
            
            # Skip empty lines
            if not text:
                i += 1
                continue
            
            # Context management (Unit, Chapter, Topic)
            if text.lower().startswith('unit-'):
                unit_name = text.split('-', 1)[1].strip()
                current_unit, _ = Unit.objects.get_or_create(
                    subject=subject, 
                    name=unit_name
                )
                current_chapter = None
                current_topic = None
                i += 1
                continue
            
            elif text.lower().startswith('chapter-'):
                chapter_name = text.split('-', 1)[1].strip()
                
                if not current_unit:
                    logger.warning(f"Chapter {chapter_name} found without a preceding unit. Skipping.")
                    i += 1
                    continue
                
                current_chapter, _ = Chapter.objects.get_or_create(
                    unit=current_unit, 
                    name=chapter_name
                )
                
                current_topic, _ = Topic.objects.get_or_create(
                    chapter=current_chapter, 
                    name=chapter_name
                )
                i += 1
                continue
            
            elif text.lower().startswith('topic-'):
                topic_name = text.split('-', 1)[1].strip()
                
                if not current_chapter:
                    logger.warning(f"Topic {topic_name} found without a preceding chapter. Skipping.")
                    i += 1
                    continue
                
                current_topic, _ = Topic.objects.get_or_create(
                    chapter=current_chapter, 
                    name=topic_name
                )
                i += 1
                continue
            
            # MCQ Detection
            if '|' in text:
                # Ensure chapter exists
                if not current_chapter:
                    logger.warning("MCQ found without a chapter. Skipping.")
                    i += 1
                    continue
                
                # If no topic was explicitly defined, use chapter name as topic
                if not current_topic:
                    current_topic, _ = Topic.objects.get_or_create(
                        chapter=current_chapter, 
                        name=current_chapter.name
                    )
                
                # Collect consecutive paragraphs for this MCQ
                mcq_paragraphs = [current_para]
                for j in range(i + 1, len(doc.paragraphs)):
                    next_para = doc.paragraphs[j]
                    # Stop if we hit another clear MCQ or context change
                    if '|' in next_para.text or len(mcq_paragraphs) >= 6:
                        break
                    mcq_paragraphs.append(next_para)
                
                # Parse the MCQ
                parts = parse_mcq(mcq_paragraphs)
                
                # Validate MCQ parts
                if not parts or len(parts) < 8:
                    logger.error(f"Incomplete MCQ data. Parts found: {parts}")
                    i += len(mcq_paragraphs)
                    continue
                
                # Validate difficulty
                difficulty = validate_difficulty(parts[7])
                if not difficulty:
                    logger.error(f"Skipping MCQ due to invalid difficulty: {parts[7]}")
                    i += len(mcq_paragraphs)
                    continue
                
                # Validate MCQ type
                mcq_type = validate_mcq_type(parts[8] if len(parts) > 8 else 'General')
                if not mcq_type:
                    logger.error(f"Skipping MCQ due to invalid type: {parts[8] if len(parts) > 8 else 'General'}")
                    i += len(mcq_paragraphs)
                    continue
                
                # Prepare MCQ data
                try:
                    mcq_data = {
                        'text': parts[0],
                        'option_1': parts[1],
                        'option_2': parts[2],
                        'option_3': parts[3],
                        'option_4': parts[4],
                        'correct_option': parts[5],
                        'explanation': parts[6],
                        'topic': current_topic,
                        'difficulty': difficulty,
                        'types': mcq_type
                    }
                    
                    # Create MCQ
                    MCQ.objects.create(**mcq_data)
                    logger.info(f"Successfully uploaded MCQ: {mcq_data['text'][:50]}...")
                    
                    # Mark paragraphs for removal
                    uploaded_mcq_indices.extend(range(i, i + len(mcq_paragraphs)))
                
                except Exception as mcq_error:
                    logger.error(f"Error creating MCQ: {mcq_error}")
                    logger.error(traceback.format_exc())
                
                # Move index to next unprocessed paragraph
                i += len(mcq_paragraphs)
                continue
            
            i += 1
        
        # Remove successfully uploaded MCQs from the document
        if uploaded_mcq_indices:
            # Create a new document
            new_doc = docx.Document()
            
            # Copy non-MCQ paragraphs and paragraphs with failed MCQs
            for i, para in enumerate(doc.paragraphs):
                if i not in uploaded_mcq_indices:
                    new_para = new_doc.add_paragraph(para.text)
                    # Copy paragraph formatting
                    new_para.style = para.style
            
            # Save the updated document
            new_doc.save(file_path)
            logger.info(f"Removed successfully uploaded MCQs from {file_path}")
        
        logger.info(f"Completed processing file: {file_path}")
    
    except Exception as file_error:
        logger.error(f"Error processing file {file_path}: {file_error}")
        logger.error(traceback.format_exc())

def main():
    """
    Main function to process MCQ files
    """
    mcq_directory = os.getcwd()
    # Directory containing MCQ DOCX files
    
    # Process all DOCX files in the directory
    for filename in os.listdir(mcq_directory):
        if filename.endswith('.docx'):
            file_path = os.path.join(mcq_directory, filename)
            process_docx_file(file_path)

if __name__ == '__main__':
    # Install required libraries
    try:
        import docx
    except ImportError:
        print("Installing required libraries...")
        import subprocess
        subprocess.check_call(['pip', 'install', 'python-docx', 'django'])
    
    main()
    print("MCQ upload process completed. Check mcq_upload_log.txt for details.")